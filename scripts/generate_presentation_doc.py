"""Generate DojoChain presentation Word document."""
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "DojoChain-Presentation-Guide.docx"


def set_doc_font(doc: Document, name: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_doc_font(doc)

    title = doc.add_heading("DojoChain 项目 Presentation 完整指南", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p(doc, "适用对象：下周团队演示与 Q&A")
    p(doc, "技术栈：Nuxt 4 + Tailwind CSS v4 + Vue 3 组件化")
    doc.add_paragraph()

    # --- 1 ---
    h(doc, "一、项目概述", 1)
    p(doc, "本项目是对 ThemeForest 官方 DojoChain WordPress 主题（dojochain.fantasythemes.net）的前端复刻，使用现代 Web 技术重建，而非 WordPress/Elementor。")
    table(doc, ["路由", "页面说明"], [
        ["/", "首页 Home — Hero、Home Demos、Features、FAQ 等"],
        ["/ico", "ICO Landing — 倒计时、Music Platform、Benefits、Team 等"],
        ["/coming-soon", "Coming Soon — 全屏倒计时 + Three.js 星空"],
        ["/privacy", "Privacy Policy — 隐私政策页"],
    ])

    # --- 2 Nuxt ---
    h(doc, "二、为什么使用 Nuxt？", 1)
    for item in [
        "文件即路由：app/pages/index.vue → /，无需手写 router",
        "组件自动注册：app/components/ 下组件可直接在模板使用",
        "SSR 友好：利于 SEO 和首屏加载",
        "Composables 复用逻辑：动画、过渡、滚动检测与 UI 分离",
        "团队协作清晰：页面薄、组件厚，职责分明",
    ]:
        bullet(doc, item)
    p(doc, "演示话术：「Nuxt 是 Vue 生态里最适合做 marketing site 的框架，路由、组件、逻辑复用都内置，比从零搭 Vite + Vue Router 更高效。」")

    # --- 3 Tailwind ---
    h(doc, "三、为什么使用 Tailwind CSS v4？", 1)
    for item in [
        "样式写在组件 class 里，改 UI 不用跳多个 CSS 文件",
        "@theme 集中管理品牌色（dojo-orange、ico-red、dojo-blue 等）",
        "响应式简单：min-[1270px]:、max-[991px]: 精确对齐原版断点",
        "Arbitrary values 支持像素级还原（如 bottom-[60px]）",
        "90%+ 样式为 utility class；main.css 仅保留 @theme、keyframes、reveal 变体",
    ]:
        bullet(doc, item)
    p(doc, "演示话术：「Tailwind 让我们快速还原设计稿，同时通过 @theme 保持全站颜色一致。」")

    # --- 4 Structure ---
    h(doc, "四、项目目录结构", 1)
    code(doc, """onboarding-task/
├── app/
│   ├── app.vue              # 全局：过渡、光标、滚动条、Cookie
│   ├── assets/css/main.css  # Tailwind + 主题色 + 动画
│   ├── components/          # UI 组件（自动注册）
│   ├── composables/         # 可复用逻辑
│   ├── pages/               # 路由页面
│   └── utils/               # 文案、导航、动画工具
├── public/images/           # 静态图片
├── nuxt.config.ts
└── package.json""")

    # --- 5 Components ---
    h(doc, "五、组件清单与目的", 1)
    h(doc, "5.1 全局组件", 2)
    table(doc, ["组件", "目的"], [
        ["DojoHeader", "顶栏 Logo、导航、购物车、手机侧栏菜单"],
        ["DojoFooter", "页脚链接、Newsletter、社交图标"],
        ["DojoPageTransition", "首次加载/切页：圆圈 + 红幕过渡"],
        ["DojoCursor", "首页自定义红色圆点光标（桌面）"],
        ["DojoScrollChrome", "Scroll 提示 + Go Top 进度条"],
        ["DojoCookieBanner", "Cookie 提示 + Privacy 链接"],
        ["DojoNavDropdown / DojoNavMobileSection", "桌面下拉 / 手机菜单"],
    ])
    h(doc, "5.2 首页 home/", 2)
    table(doc, ["组件", "目的"], [
        ["HomeHero", "Hero 背景、标题、按钮、四个特性"],
        ["HomeSections", "Demos、Features、介绍区、FAQ"],
        ["HomeDemoCard", "Demo 卡片：从右滑入 + hover 缓慢放大"],
        ["HomeBtn / HomeHeroIcon / HomeFeatureIcon", "按钮与图标"],
    ])
    h(doc, "5.3 ICO 页 ico/", 2)
    table(doc, ["组件", "目的"], [
        ["IcoHero", "ICO 倒计时、标题、Try Demo 按钮"],
        ["IcoSections", "Music Revolution、Benefits、Team 等"],
        ["IcoBtn", "带右侧图标的 pill 按钮"],
        ["IcoProgressCircle", "Token 分配环形进度动画"],
    ])
    h(doc, "5.4 Coming Soon", 2)
    table(doc, ["组件", "目的"], [
        ["ComingSoonHero", "倒计时与文案"],
        ["DojoSkyBg", "Three.js 星空背景"],
    ])
    p(doc, "拆组件的原因：复用（IcoBtn）、易维护（改 Hero 只改一个文件）、页面结构清晰、与原版 section 一一对应。")

    # --- 6 Usage ---
    h(doc, "六、如何使用与修改组件", 1)
    h(doc, "6.1 在页面中组合", 2)
    code(doc, """<!-- app/pages/index.vue -->
<DojoHeader variant="default" active-nav="home" />
<HomeHero />
<HomeSections />""")
    h(doc, "6.2 传 Props", 2)
    code(doc, """<DojoHeader variant="ico" active-nav="landings" />
<HomeDemoCard title="ICO" image="/images/home/ico-small.jpg" href="/ico" />
<IcoBtn variant="primary" icon="fas fa-music">Try Demo</IcoBtn>""")
    h(doc, "6.3 修改文案/数据", 2)
    bullet(doc, "首页 demo 列表、FAQ：app/utils/homeContent.ts")
    bullet(doc, "导航菜单：app/utils/navigation.ts")
    h(doc, "6.4 修改品牌色", 2)
    bullet(doc, "编辑 app/assets/css/main.css 的 @theme 中 --color-dojo-orange 等")
    bullet(doc, "全站 bg-dojo-orange、text-ico-red 等 class 会同步更新")

    # --- 7 Animations ---
    h(doc, "七、动画系统说明", 1)
    h(doc, "7.1 页面预加载（DojoPageTransition）", 2)
    bullet(doc, "流程：白圈放大 → 红幕从底部滑上 → 显示页面")
    bullet(doc, "时间：app/composables/usePageTransition.ts（CIRCLE_DONE_MS、CURTAIN_MS）")
    bullet(doc, "曲线：main.css 中 @keyframes preloader-curtain-sweep")
    h(doc, "7.2 滚动 Reveal（滚到才出现）", 2)
    bullet(doc, "工具：app/utils/reveal.ts — revealUp / revealRight / revealLeft")
    bullet(doc, "逻辑：app/composables/useIcoScrollReveal.ts — IntersectionObserver")
    bullet(doc, "Keyframes：main.css — ico-reveal-from-up 等")
    code(doc, """<!-- 单个元素 -->
<div :class="revealUp" :style="{ '--reveal-delay': `${i * 200}ms` }">

<!-- 整组触发 -->
<section data-reveal-trigger="features">
  <div data-reveal-group="features" :class="revealUp">...</div>
</section>

<!-- 页面启用 -->
<div ref="pageRoot" :data-anim-ready="animReady || undefined">
const { animReady } = useIcoScrollReveal(pageRoot)""")
    h(doc, "7.3 Home Demos 卡片（HomeDemoCard.vue）", 2)
    bullet(doc, "进入视口：黑色遮罩揭开 + 图片从右滑入")
    bullet(doc, "Hover：1.2s 缓慢放大至 scale(1.05)")
    bullet(doc, "可调：translate-x-[150px]、duration-[2000ms]、group-hover:scale-105")
    h(doc, "7.4 按钮与微交互", 2)
    bullet(doc, "IcoBtn：hover 缓慢放大 + 图标上下滑动")
    bullet(doc, "DojoCursor：hover 链接时红点放大")

    # --- 8 Mobile ---
    h(doc, "八、Mobile 响应式如何保证", 1)
    table(doc, ["断点", "用途"], [
        ["默认", "手机优先布局"],
        ["max-[991px]:", "ICO Hero 倒计时 2 列、间距调整"],
        ["min-[1270px]:", "完整桌面横排导航"],
        ["md: / lg: / xl:", "栅格列数（Features 1→2→3 列）"],
    ])
    bullet(doc, "<1270px：汉堡菜单 + 右侧滑出 side menu")
    bullet(doc, "Scroll 提示 / Go Top：max-[992px]:hidden（仅桌面）")
    bullet(doc, "自定义光标：仅 pointer:fine 设备 + 特定路由")
    p(doc, "测试：npm run dev → Chrome DevTools 设备模式 → 测 iPhone/iPad 宽度")

    # --- 9 Assets ---
    h(doc, "九、图片与 Logo 存放与引用", 1)
    p(doc, "静态资源放在 public/ 目录，浏览器通过根路径访问。")
    code(doc, """public/images/logo-light.svg     → /images/logo-light.svg
public/images/home/hero-bg.jpg   → /images/home/hero-bg.jpg""")
    p(doc, "引用方式：")
    code(doc, """<img src="/images/home/ico-small.jpg" alt="ICO" />
class="bg-[url('/images/home/hero-bg.jpg')] bg-cover"
homeContent.ts: { image: '/images/home/ico-small.jpg', href: '/ico' }""")
    p(doc, "换图步骤：1) 文件放入 public/images/  2) 改组件 src 或 homeContent.ts  3) 刷新浏览器")
    p(doc, "注意：部分 ICO 图片（Team 照片、Featured 品牌 logo）仍引用官方 CDN，上线前应全部本地化。")

    # --- 10 Run ---
    h(doc, "十、如何运行项目", 1)
    code(doc, """cd onboarding-task
npm install
npm run dev
# 浏览器打开 http://localhost:3000

npm run build    # 生产构建
npm run preview  # 预览构建结果""")
    p(doc, "环境要求：Node.js 18+（建议 20+）")

    # --- 11 Presentation flow ---
    h(doc, "十一、建议演示流程（5–10 分钟）", 1)
    steps = [
        "30 秒：说明项目目标 — Nuxt + Tailwind 复刻 DojoChain",
        "1 分钟：为什么选 Nuxt + Tailwind",
        "2 分钟：Live demo — / → 滚到 Demos → 点 ICO → /ico",
        "2 分钟：打开 index.vue + HomeHero.vue 讲组件结构",
        "1 分钟：讲 reveal 动画系统",
        "1 分钟：public/images/ 与换资源",
        "Q&A",
    ]
    for i, s in enumerate(steps, 1):
        bullet(doc, s)

    # --- 12 Q&A ---
    h(doc, "十二、常见问题 Q&A", 1)
    qa = [
        ("为什么不用 WordPress 原版？", "原版依赖 PHP + Elementor。Nuxt 提供现代前端栈、更好性能、可部署到 Vercel/Netlify，不绑 WordPress。"),
        ("为什么不用 Bootstrap/MUI？", "Marketing landing 需要像素级 custom 设计，Tailwind 更快还原，Bootstrap 组件风格固定需大量 override。"),
        ("样式真的全是 Tailwind 吗？", "90%+ 是 utility class。main.css 仅 @theme、keyframes、reveal 变体，属 Tailwind v4 标准做法。"),
        ("怎么加新页面？", "在 app/pages/ 新建 about.vue → 自动 /about，组合 Header + 内容 + Footer。"),
        ("Mobile 和 Desktop 是两套 HTML 吗？", "不是，同一套 HTML，Tailwind breakpoint 切换布局。"),
        ("自定义光标为什么 Privacy 没有？", "故意设计：阅读页用系统光标，体验更正常。"),
        ("红幕预加载是什么？", "模仿原版切页效果；红幕仅在动画开始时渲染，避免加载前底部露红。"),
        ("build 失败怎么办？", "看终端报错；常见为路径/组件名错误；可删 node_modules 重装。"),
        ("和原版有多像？", "布局、颜色、字体、主要 section、动画行为对照官方 demo；Shop/Blog 等子站未全做。"),
        ("composables 是什么？", "Vue 3 可复用函数，类似 React hooks，封装 scroll reveal 等逻辑。"),
    ]
    for q, a in qa:
        p(doc, f"Q: {q}", bold=True)
        p(doc, f"A: {a}")
        doc.add_paragraph()

    # --- 13 Improvements ---
    h(doc, "十三、后续可改进方向（Improvements）", 1)
    p(doc, "以下按优先级排列，可在 presentation 中作为「未来工作」或 Q&A 补充。")

    h(doc, "13.1 功能完整性", 2)
    for item in [
        "补全其余 Landing 页：Crypto Wallet、Fintech、Crypto News、Shop（导航已有，路由未实现）",
        "将 href=\"#\" 占位链接改为真实路由或外链（Header 社交、Footer 链接、Whitepaper 等）",
        "实现搜索功能（Header 搜索框目前为 UI）",
        "Newsletter 表单接入后端或第三方（Mailchimp 等）",
        "Shop / Cart 页面（原版有 WooCommerce demo）",
    ]:
        bullet(doc, item)

    h(doc, "13.2 资源与性能", 2)
    for item in [
        "将所有 CDN 图片下载到 public/images/（ICO Team、Featured 品牌、about/technology 大图）",
        "确认 logo-light.svg、ico-hero-bg.jpg、coming-soon-logo.png 等文件在 public 中齐全",
        "图片 WebP/AVIF 格式 + lazy loading 优化（部分已有 loading=\"lazy\"）",
        "Three.js 星空按需加载，Coming Soon 页外不加载 three 包",
        "Font Awesome 改为按需 SVG 图标，减少全量 CSS 请求",
    ]:
        bullet(doc, item)

    h(doc, "13.3 质量与可访问性", 2)
    for item in [
        "补充 E2E 测试（项目已有 Playwright devDependency，可写首页/ICO 关键路径测试）",
        "Accessibility：键盘导航、focus 状态、aria-label 审查",
        "SEO：每页独立 meta description、Open Graph 图片",
        "404 自定义页面（导航 Inner Pages 提到 Error 404）",
        "i18n 多语言支持（若目标用户非英语区）",
    ]:
        bullet(doc, item)

    h(doc, "13.4 代码与架构", 2)
    for item in [
        "README 更新为 DojoChain 专用说明（当前仍为 Nuxt starter 模板）",
        "删除 repo 根目录参考 HTML 文件（_home-page.html 等），减少混淆",
        "Content 可迁到 CMS（Sanity/Contentful）或 JSON，非开发者也能改文案",
        "统一 IcoBtn 与 HomeBtn 为通用 DojoBtn 组件（减少重复）",
        "TypeScript 严格模式与 props 类型补全",
    ]:
        bullet(doc, item)

    h(doc, "13.5 UX 细节", 2)
    for item in [
        "页面切换过渡：更多路由使用 navigateWithTransition",
        "Home Demos 其余卡片链到对应 landing（目前多为 #）",
        "ICO 倒计时可改为真实目标日期配置化",
        "Cookie banner 持久化（localStorage 记住用户选择）",
        "Dark/light scroll 指示器 tone 检测可进一步微调",
    ]:
        bullet(doc, item)

    h(doc, "13.6 部署与 DevOps", 2)
    for item in [
        "CI/CD：GitHub Actions 自动 build + lint",
        "部署到 Vercel / Netlify / Azure Static Web Apps",
        "环境变量管理（API keys、CMS token）",
        "Lighthouse 性能/SEO 分数 baseline 与监控",
    ]:
        bullet(doc, item)

    # --- 14 Summary ---
    h(doc, "十四、一句话总结", 1)
    p(doc, "本项目用 Nuxt 4 管理路由与组件架构，用 Tailwind v4 还原 DojoChain 设计；页面拆成可复用 component，动画用 IntersectionObserver + CSS keyframes；静态资源放 public/images/；通过 responsive Tailwind 保证 mobile/desktop。后续重点是补全页面、本地化资源、测试与部署。")

    doc.add_paragraph()
    p(doc, "文档生成位置：项目根目录 DojoChain-Presentation-Guide.docx")
    p(doc, "祝演示顺利！")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
